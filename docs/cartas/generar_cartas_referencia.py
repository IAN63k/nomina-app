#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de cartas de vacaciones (combinación de correspondencia / mail merge).

Toma una plantilla de Word (.docx) que contiene marcadores entre llaves, por
ejemplo {NOMBRE} o {DIAS_TIENE}, y un archivo de Excel cuya fila de encabezados
usa exactamente esos mismos nombres. Por cada fila con datos del Excel genera
una carta nueva reemplazando los marcadores por los valores de esa fila.

Uso básico (con los nombres por defecto):
    python generar_cartas.py

Uso indicando los archivos:
    python generar_cartas.py --plantilla CARTA.docx --excel datos.xlsx --hoja PROGRAMACION

El formato (negritas, fuentes, tamaños, etc.) de la plantilla se conserva,
porque solo se reemplaza el texto dentro de cada "run" sin tocar su estilo.

Requisitos:
    pip install python-docx openpyxl
"""

import argparse
import datetime
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.exit("Falta la librería python-docx. Instálala con:  pip install python-docx")

try:
    import openpyxl
except ImportError:
    sys.exit("Falta la librería openpyxl. Instálala con:  pip install openpyxl")


# --- Configuración por defecto -------------------------------------------------

PLANTILLA_DEFECTO = "CARTA_DE_VACACIONES_JUNIO_2026.docx"
EXCEL_DEFECTO = "datos.xlsx"
HOJA_DEFECTO = "PROGRAMACION"   # hoja del Excel cuyos encabezados coinciden con la carta
CARPETA_SALIDA = "cartas_generadas"

# Meses en español para formatear las fechas de forma legible.
MESES_ES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

# Columnas de respaldo: si la columna de la izquierda está vacía en una fila,
# el programa toma el valor de la columna de la derecha. Se deja vacío por
# defecto: las fechas ya vienen resueltas en el Excel.
RESPALDOS = {}

# Columnas cuyas fechas se muestran en formato corto dd/mm/aaaa (en lugar del
# formato largo en español "1 de junio de 2026"). Las fechas de contrato van
# en formato corto; las de vacaciones quedan en formato largo.
COLUMNAS_FECHA_CORTA = {"PERIODO 1", "PERIODO 2"}


# --- Utilidades ----------------------------------------------------------------

def formatear_valor(valor, fecha_corta=False):
    """Convierte un valor de celda de Excel en texto apto para la carta.

    Si fecha_corta es True, las fechas se devuelven como dd/mm/aaaa.
    """
    if valor is None:
        return ""

    # Fechas
    if isinstance(valor, (datetime.datetime, datetime.date)):
        # Las fechas de contrato pueden venir como 30/12/1899 (sin vencimiento);
        # en formato corto SÍ se muestran. En formato largo, una fecha tan
        # antigua se trata como vacía.
        if fecha_corta:
            return f"{valor.day:02d}/{valor.month:02d}/{valor.year}"
        if valor.year < 1900:
            return ""
        return f"{valor.day} de {MESES_ES[valor.month]} de {valor.year}"

    # Una hora suelta (p. ej. 00:00) significa celda sin fecha real.
    # En las columnas de fecha corta, ese caso corresponde a la fecha origen
    # de Excel (30/12/1899) que se usa para contratos sin vencimiento.
    if isinstance(valor, datetime.time):
        if fecha_corta:
            return "30/12/1899"
        return ""

    # Números enteros sin decimales -> sin ".0"
    if isinstance(valor, float) and valor.is_integer():
        return str(int(valor))

    return str(valor).strip()


def esta_vacio(valor):
    """Indica si un valor de celda debe considerarse vacío.

    Una hora suelta (time) o un datetime con año anterior a 1900 se tratan
    como vacíos, igual que None o cadenas en blanco. Esto detecta las celdas
    de fecha que Excel guarda como 00:00 cuando en realidad están vacías.
    """
    if valor is None:
        return True
    if isinstance(valor, datetime.time):
        return True
    if isinstance(valor, datetime.datetime) and valor.year < 1900:
        return True
    if isinstance(valor, str) and valor.strip() == "":
        return True
    return False


def limpiar_nombre_archivo(texto):
    """Deja un texto seguro para usarlo como nombre de archivo."""
    texto = re.sub(r'[\\/:*?"<>|]', "_", str(texto)).strip()
    texto = re.sub(r"\s+", "_", texto)
    return texto[:120] if texto else "sin_nombre"


def reemplazar_en_parrafo(parrafo, reemplazos):
    """Reemplaza los marcadores dentro de un párrafo conservando el formato.

    Trabaja run por run. Si algún marcador quedara partido entre varios runs,
    como respaldo une el texto del párrafo y lo vuelve a escribir en el primer run.
    """
    # 1) Intento normal: cada run por separado (preserva todos los estilos).
    hubo_marcador_partido = False
    for run in parrafo.runs:
        texto = run.text
        if "{" in texto:
            nuevo = aplicar_reemplazos(texto, reemplazos)
            run.text = nuevo
            # ¿quedó una llave suelta? -> el marcador estaba partido
            if "{" in nuevo and "}" not in nuevo:
                hubo_marcador_partido = True

    # 2) Respaldo: si detectamos marcadores partidos entre runs, reconstruimos.
    texto_completo = "".join(r.text for r in parrafo.runs)
    if hubo_marcador_partido or ("{" in texto_completo and any(m in texto_completo for m in reemplazos)):
        nuevo_completo = aplicar_reemplazos(texto_completo, reemplazos)
        if nuevo_completo != texto_completo and parrafo.runs:
            parrafo.runs[0].text = nuevo_completo
            for run in parrafo.runs[1:]:
                run.text = ""


def aplicar_reemplazos(texto, reemplazos):
    """Sustituye todos los marcadores {CLAVE} presentes en un texto."""
    for marcador, valor in reemplazos.items():
        if marcador in texto:
            texto = texto.replace(marcador, valor)
    return texto


def iterar_parrafos(doc):
    """Devuelve todos los párrafos del documento, incluidos los de tablas,
    encabezados y pies de página."""
    for p in doc.paragraphs:
        yield p
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    yield p
    for seccion in doc.sections:
        for contenedor in (seccion.header, seccion.footer,
                           seccion.first_page_header, seccion.first_page_footer):
            for p in contenedor.paragraphs:
                yield p


# --- Lógica principal ----------------------------------------------------------

def leer_filas_excel(ruta_excel, nombre_hoja):
    """Lee la hoja indicada y devuelve (encabezados, lista_de_filas_dict)."""
    wb = openpyxl.load_workbook(ruta_excel, data_only=True)

    if nombre_hoja not in wb.sheetnames:
        sys.exit(
            f"No existe la hoja '{nombre_hoja}' en el Excel.\n"
            f"Hojas disponibles: {', '.join(wb.sheetnames)}"
        )

    ws = wb[nombre_hoja]

    # Encabezados (fila 1)
    encabezados = []
    for celda in ws[1]:
        encabezados.append(str(celda.value).strip() if celda.value is not None else None)

    filas = []
    for fila in ws.iter_rows(min_row=2, values_only=True):
        registro = {}
        tiene_datos = False
        for col, valor in zip(encabezados, fila):
            if col is None:
                continue
            registro[col] = valor
            if valor not in (None, ""):
                tiene_datos = True
        if tiene_datos:
            filas.append(registro)

    return [e for e in encabezados if e], filas


def generar(plantilla, excel, hoja, carpeta_salida):
    ruta_plantilla = Path(plantilla)
    ruta_excel = Path(excel)

    if not ruta_plantilla.exists():
        sys.exit(f"No se encontró la plantilla: {ruta_plantilla}")
    if not ruta_excel.exists():
        sys.exit(f"No se encontró el Excel: {ruta_excel}")

    encabezados, filas = leer_filas_excel(ruta_excel, hoja)
    print(f"Hoja '{hoja}': {len(filas)} filas con datos.")
    print(f"Columnas detectadas: {', '.join(encabezados)}")

    # Avisar qué marcadores de la plantilla no tienen columna en el Excel.
    doc_muestra = Document(str(ruta_plantilla))
    texto_plantilla = "\n".join(p.text for p in iterar_parrafos(doc_muestra))
    marcadores_plantilla = set(re.findall(r"\{[^}]+\}", texto_plantilla))
    sin_columna = [m for m in marcadores_plantilla if m.strip("{}") not in encabezados]
    if sin_columna:
        print("\n[Aviso] Estos marcadores de la carta NO tienen una columna con el "
              "mismo nombre en el Excel y quedarán vacíos o sin cambiar:")
        for m in sorted(sin_columna):
            print(f"   {m}")
    print()

    salida = Path(carpeta_salida)
    salida.mkdir(parents=True, exist_ok=True)

    generadas = 0
    usos_respaldo = {origen: 0 for origen in RESPALDOS}
    for indice, registro in enumerate(filas, start=1):
        # Aplicar respaldos: si una columna está vacía en esta fila, usar el
        # valor de su columna de respaldo (p. ej. FECHA FIN -> PERIODO 2).
        for col_destino, col_origen in RESPALDOS.items():
            if col_destino in registro and esta_vacio(registro.get(col_destino)):
                valor_respaldo = registro.get(col_origen)
                if not esta_vacio(valor_respaldo):
                    registro[col_destino] = valor_respaldo
                    usos_respaldo[col_destino] += 1

        # Construir el diccionario de reemplazos: {COLUMNA} -> valor formateado
        reemplazos = {
            "{" + columna + "}": formatear_valor(
                valor, fecha_corta=(columna in COLUMNAS_FECHA_CORTA)
            )
            for columna, valor in registro.items()
        }

        # Abrir una copia limpia de la plantilla por cada carta.
        doc = Document(str(ruta_plantilla))
        for parrafo in iterar_parrafos(doc):
            if "{" in parrafo.text:
                reemplazar_en_parrafo(parrafo, reemplazos)

        # Nombre de archivo: usa NOMBRE + APELLIDO o CEDULA si existen.
        partes = []
        for clave in ("NOMBRE", "APELLIDO"):
            if registro.get(clave):
                partes.append(formatear_valor(registro[clave]))
        etiqueta = " ".join(partes) if partes else f"fila_{indice}"

        nombre = f"{limpiar_nombre_archivo(etiqueta)}.docx"
        doc.save(str(salida / nombre))
        generadas += 1
        print(f"  [{indice:>3}/{len(filas)}] {nombre}")

    print(f"\nListo. Se generaron {generadas} cartas en la carpeta: {salida.resolve()}")

    respaldos_usados = {c: n for c, n in usos_respaldo.items() if n}
    if respaldos_usados:
        print("\n[Info] Como algunas celdas estaban vacías, se usaron columnas de respaldo:")
        for col_destino, n in respaldos_usados.items():
            print(f"   {col_destino}: se tomó '{RESPALDOS[col_destino]}' en {n} fila(s).")


def main():
    parser = argparse.ArgumentParser(
        description="Genera una carta de Word por cada fila de un Excel, "
                    "reemplazando los marcadores {ENTRE_LLAVES} de la plantilla."
    )
    parser.add_argument("--plantilla", default=PLANTILLA_DEFECTO,
                        help=f"Ruta de la plantilla .docx (por defecto: {PLANTILLA_DEFECTO})")
    parser.add_argument("--excel", default=EXCEL_DEFECTO,
                        help=f"Ruta del Excel con los datos (por defecto: {EXCEL_DEFECTO})")
    parser.add_argument("--hoja", default=HOJA_DEFECTO,
                        help=f"Nombre de la hoja del Excel (por defecto: {HOJA_DEFECTO})")
    parser.add_argument("--salida", default=CARPETA_SALIDA,
                        help=f"Carpeta donde guardar las cartas (por defecto: {CARPETA_SALIDA})")
    args = parser.parse_args()

    generar(args.plantilla, args.excel, args.hoja, args.salida)


if __name__ == "__main__":
    main()
